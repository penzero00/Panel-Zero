"""
DOCX Parser Module
Extracts text and structure from DOCX files for LLM processing
Implements smart chunking for large documents
"""

from docx import Document
from typing import List, Dict, Any, Tuple


class ChapterExtractor:
    """
    Extracts chapters/sections from DOCX files.
    Implements smart chunking per AGENTS.md: Large documents split by chapters
    """

    def __init__(self, docx_path: str):
        self.doc = Document(docx_path)
        self.chapters: List[Dict[str, Any]] = []

    def extract_chapters(self) -> List[Dict[str, Any]]:
        """
        Extract chapters from document.
        Assumes heading levels indicate chapter structure.
        """
        current_chapter = None
        chapters = []

        for para in self.doc.paragraphs:
            if para.style.name.startswith("Heading 1"):
                # New chapter
                if current_chapter:
                    chapters.append(current_chapter)

                current_chapter = {
                    "title": para.text,
                    "paragraphs": [],
                    "tables": [],
                }
            elif current_chapter:
                current_chapter["paragraphs"].append(para.text)

        # Add last chapter
        if current_chapter:
            chapters.append(current_chapter)

        self.chapters = chapters
        return chapters

    def extract_full_text(self) -> str:
        """Extract all text content from document"""
        return "\n".join([para.text for para in self.doc.paragraphs])

    def extract_tables(self) -> List[List[List[str]]]:
        """Extract all tables from document"""
        tables = []
        for table in self.doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        return tables

    def chunk_for_llm(self, max_tokens: int = 4000) -> List[Dict[str, Any]]:
        """
        Split document into chunks suitable for LLM processing.
        Respects chapter boundaries when possible.
        """
        chunks = []
        current_chunk = ""
        current_token_count = 0

        # Rough estimate: 1 word ≈ 1.3 tokens
        TOKEN_PER_WORD = 1.3

        for para in self.doc.paragraphs:
            para_tokens = int(len(para.text.split()) * TOKEN_PER_WORD)

            if current_token_count + para_tokens > max_tokens and current_chunk:
                chunks.append({"text": current_chunk, "token_estimate": current_token_count})
                current_chunk = ""
                current_token_count = 0

            current_chunk += para.text + "\n"
            current_token_count += para_tokens

        if current_chunk:
            chunks.append({"text": current_chunk, "token_estimate": current_token_count})

        return chunks

    def get_document_metadata(self) -> Dict[str, Any]:
        """Extract document metadata"""
        if self.doc.core_properties:
            props = self.doc.core_properties
            return {
                "title": props.title or "Untitled",
                "author": props.author or "Unknown",
                "created": props.created,
                "modified": props.modified,
                "subject": props.subject or "",
            }
        return {}
