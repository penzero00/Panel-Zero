"""
Surgical DOCX Injection Module
Implements non-destructive XML-based editing for DOCX files
Following AGENTS.md: Golden Rules for Document Processing
"""

import os
import shutil
from pathlib import Path
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from typing import Optional, List, Dict, Any

class SurgicalInjector:
    """
    Handles non-destructive DOCX editing via XML manipulation.
    CRITICAL: Always copy the original file before processing.
    CRITICAL: Apply highlights directly to run objects to preserve styles.
    """

    def __init__(self, docx_path: str):
        """
        Initialize with DOCX file path.
        Automatically creates a working copy so the original downloaded from Supabase is safe.
        """
        self.original_path = docx_path
        
        # Create a safe working copy path (e.g., /tmp/file_working.docx)
        path_obj = Path(docx_path)
        self.working_path = str(path_obj.with_name(f"{path_obj.stem}_working{path_obj.suffix}"))
        
        # Create the working copy
        shutil.copy2(self.original_path, self.working_path)
        print(f"Working copy created at: {self.working_path}")
        
        # Open the working copy for manipulation
        self.doc = Document(self.working_path)

    def inject_issue_highlights(self, issues: List[Dict[str, Any]], max_highlights: int = 10000) -> Dict[str, Any]:
        """
        Apply highlights and comments based on analysis issues.
        Uses location.text field to find exact text to highlight within paragraphs.
        Implements advanced fuzzy matching for better highlighting success rate.
        Major severity = Red highlight, Minor severity = Yellow highlight
        No practical limit on highlights (default 10000 is very high).
        """
        if not self.doc:
            return {"success": False, "error": "No document loaded"}
        
        try:
            from difflib import SequenceMatcher
            
            successful_highlights = 0
            skipped_issues = 0
            not_found_issues = 0
            
            # Process all issues up to max_highlights
            issues_to_process = issues[:max_highlights]
            if len(issues) > max_highlights:
                print(f"[SURGICAL_INJECTOR] Processing {max_highlights} of {len(issues)} issues")
            else:
                print(f"[SURGICAL_INJECTOR] Processing all {len(issues)} issues")
            
            # Build full document text map for better searching
            full_doc_text = "\n".join([p.text for p in self.doc.paragraphs])
            
            for issue_idx, issue in enumerate(issues_to_process):
                severity = issue.get("severity", "minor")
                
                # Determine highlight color
                if severity and str(severity).lower() in ["major", "high", "critical"]:
                    highlight_color = WD_COLOR_INDEX.RED
                else:
                    highlight_color = WD_COLOR_INDEX.YELLOW
                
                # Try to find text to highlight - with multiple fallback strategies
                text_to_highlight = None
                
                # Strategy 1: Try location.text field
                if "location" in issue and isinstance(issue["location"], dict):
                    text_to_highlight = issue["location"].get("text")
                
                # Strategy 2: Try text field
                if not text_to_highlight and "text" in issue:
                    text_to_highlight = issue.get("text")
                
                # Strategy 3: Try issue description (first 80 chars)
                if not text_to_highlight and "issue" in issue:
                    text_to_highlight = issue["issue"][:80]
                
                if not text_to_highlight or len(text_to_highlight.strip()) < 2:
                    skipped_issues += 1
                    continue
                
                # Search for text in document with advanced matching
                search_text = text_to_highlight.strip()
                found = False
                best_match_ratio = 0
                best_match_para = None
                best_match_run = None
                
                # Strategy 1: Exact match in paragraphs
                for para_idx, para in enumerate(self.doc.paragraphs):
                    para_text = para.text
                    
                    # Check if search text is in paragraph
                    if search_text in para_text:
                        # Highlight runs containing this text
                        for run in para.runs:
                            run_text = run.text
                            if search_text in run_text:
                                run.font.highlight_color = highlight_color
                                successful_highlights += 1
                                found = True
                                break
                        if found:
                            break
                
                # Strategy 2: Case-insensitive search
                if not found:
                    search_lower = search_text.lower()
                    for para_idx, para in enumerate(self.doc.paragraphs):
                        para_text_lower = para.text.lower()
                        
                        if search_lower in para_text_lower:
                            for run in para.runs:
                                if search_lower in run.text.lower():
                                    run.font.highlight_color = highlight_color
                                    successful_highlights += 1
                                    found = True
                                    break
                        if found:
                            break
                
                # Strategy 3: Partial word match (first 25 chars)
                if not found and len(search_text) > 5:
                    partial_search = search_text[:25].strip()
                    for para_idx, para in enumerate(self.doc.paragraphs):
                        if partial_search in para.text:
                            for run in para.runs:
                                if partial_search in run.text:
                                    run.font.highlight_color = highlight_color
                                    successful_highlights += 1
                                    found = True
                                    break
                        if found:
                            break
                
                # Strategy 4: Fuzzy matching using similarity ratio
                if not found:
                    for para_idx, para in enumerate(self.doc.paragraphs):
                        para_text = para.text
                        # Check similarity with threshold of 0.6
                        ratio = SequenceMatcher(None, search_lower, para_text.lower()).ratio()
                        if ratio > 0.6 and ratio > best_match_ratio:
                            best_match_ratio = ratio
                            best_match_para = para
                    
                    if best_match_para:
                        # Highlight the first run in the best matching paragraph
                        if best_match_para.runs:
                            run = best_match_para.runs[0]
                            run.font.highlight_color = highlight_color
                            successful_highlights += 1
                            found = True
                
                # Strategy 5: Extract key words and search for them
                if not found and len(search_text) > 10:
                    # Extract first few meaningful words
                    words = [w for w in search_text.split() if len(w) > 3][:3]
                    if words:
                        search_keywords = " ".join(words)
                        search_keywords_lower = search_keywords.lower()
                        
                        for para in self.doc.paragraphs:
                            if search_keywords_lower in para.text.lower():
                                for run in para.runs:
                                    if any(w.lower() in run.text.lower() for w in words):
                                        run.font.highlight_color = highlight_color
                                        successful_highlights += 1
                                        found = True
                                        break
                                if found:
                                    break
                
                if not found:
                    not_found_issues += 1
            
            return {
                "success": True,
                "highlights_applied": successful_highlights,
                "skipped": skipped_issues,
                "not_found": not_found_issues,
                "total_processed": len(issues_to_process),
                "total_issues": len(issues),
                "message": f"Applied {successful_highlights} highlights ({not_found_issues} text not found, {skipped_issues} skipped)"
            }
        except Exception as e:
            print(f"Error injecting highlights: {e}")
            import traceback
            traceback.print_exc()
            return {"success": False, "error": str(e)}
    
    def inject_yellow_highlight(self, paragraph_index: int, run_index: int, text_match: str = "") -> bool:
        """
        Inject yellow highlight on specific run within a paragraph.
        """
        try:
            if paragraph_index >= len(self.doc.paragraphs):
                return False

            para = self.doc.paragraphs[paragraph_index]
            
            if run_index >= len(para.runs):
                return False

            run = para.runs[run_index]

            # Verify text match for safety (optional but recommended)
            if text_match and text_match not in run.text:
                print(f"Text mismatch: Expected '{text_match}', found '{run.text}'")
                return False

            # Apply highlight without altering other styles
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            return True

        except Exception as e:
            print(f"Error injecting highlight: {e}")
            return False

    def save_processed_file(self, output_path: str) -> bool:
        """
        Save the processed document to the final output path.
        Cleans up the temporary working file afterward.
        """
        try:
            self.doc.save(output_path)
            print(f"Processed document saved to: {output_path}")
            
            # Clean up the temporary working file
            if os.path.exists(self.working_path):
                os.remove(self.working_path)
                
            return True
        except Exception as e:
            print(f"Error saving processed file: {e}")
            return False

def process_docx_with_highlights(
    input_path: str,
    output_path: str,
    highlights: List[Dict[str, Any]],
) -> Dict[str, Any]:
    """
    Main function called by the Celery Worker to process a DOCX file.
    """
    try:
        injector = SurgicalInjector(input_path)

        # Apply highlights based on AI output
        successful_highlights = 0
        for highlight in highlights:
            # AI will provide these coordinates
            if injector.inject_yellow_highlight(
                paragraph_index=highlight.get("paragraph_index", 0),
                run_index=highlight.get("run_index", 0),
                text_match=highlight.get("text_match", ""),
            ):
                successful_highlights += 1

        # Save processed file (e.g., thesis_REVIEWED.docx)
        if injector.save_processed_file(output_path):
            return {
                "success": True,
                "output_path": output_path,
                "highlights_applied": successful_highlights,
                "message": f"Successfully processed {successful_highlights} highlights",
            }
        else:
            return {"success": False, "error": "Failed to save processed file"}

    except Exception as e:
        return {"success": False, "error": str(e)}