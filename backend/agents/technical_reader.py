import os
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
from .llm_executor import LLMExecutor

class TechnicalReaderAgent(LLMExecutor):
    """
    Technical Reader Agent: Analyzes methodology, code quality, and technical correctness.
    Combines Zero-Tolerance Format Checking with AI critique and surgical injection markers.
    """
    def __init__(self, docx_path: Optional[str] = None, agent_profile: Optional[Dict[str, Any]] = None):
        super().__init__()
        
        # Structural/Formatting Setup
        self.docx_path = docx_path
        if docx_path and os.path.exists(docx_path):
            self.doc = Document(docx_path)
        else:
            self.doc = None
            
        self.agent_profile = agent_profile or {}
        self.errors: List[Dict[str, Any]] = []
        self.feedback_with_locations: List[Dict[str, Any]] = []

        # AI Configuration - use Azure model
        self.agent_model = self.default_model
        base_system_prompt = (
            "You are the Technical Reader on a thesis defense panel. "
            "Your job is to strictly evaluate the methodology, technical architecture, "
            "and logical flow of the system described in the thesis. "
            "IMPORTANT INSTRUCTION:\n"
            "1. Check ALL document formatting, fonts, margins, and styles against the PROFILE PREFERENCES below.\n"
            "2. ANY deviation from the profile preferences MUST be flagged as a formatting issue.\n"
            "3. Find ALL technical issues AND all formatting/preference violations.\n"
            "4. Report EVERY single issue found, no matter how many.\n\n"
            "Format your response as JSON with 'issues' array. Each issue must have: "
            "location (object with 'text' field containing the exact problematic text from document), "
            "type (string like 'technical|methodology|formatting|font|margin|paragraph'), "
            "severity (must be 'major' or 'minor'), "
            "issue (string description of the problem), suggestion (string with recommended fix). "
            "Include the EXACT text snippet in location.text so it can be highlighted. "
            "Find and report ALL technical and formatting issues without limits."
        )
        
        # Add comprehensive preferences from profile (including cross-domain preferences)
        comprehensive_prefs = self.build_comprehensive_preferences(self.agent_profile)
        if comprehensive_prefs:
            base_system_prompt += f"\n\n{comprehensive_prefs}\n\n"
            base_system_prompt += "===\nCRITICAL: The above PROFILE PREFERENCES are MANDATORY REQUIREMENTS.\n"
            base_system_prompt += "STRICTLY VERIFY that the entire document matches EVERY enabled preference.\n"
            base_system_prompt += "If ANY enabled preference is NOT met, flag it with severity=major.\n"
            base_system_prompt += "==="
        
        # Append custom instructions if provided in agent profile
        custom_instruction = self.agent_profile.get("custom_instruction", "").strip()
        if custom_instruction:
            self.system_prompt = f"{base_system_prompt}\n\nADDITIONAL INSTRUCTIONS FROM USER:\n{custom_instruction}\n\n"
            self.system_prompt += "REMINDER: Custom instructions SUPPLEMENT the critical preference checks above. Check both equally thoroughly."
        else:
            self.system_prompt = base_system_prompt

    def check_margins(self) -> Dict[str, Any]:
        """Check page margins against agent profile preferences"""
        if not self.doc:
            return {"success": True, "error": "No docx file provided for margin checking."}
        
        # Use profile settings if available, otherwise use defaults
        requirements = {
            "left": self.agent_profile.get("margin_left_inches") if self.agent_profile else 1.5,
            "right": self.agent_profile.get("margin_right_inches") if self.agent_profile else 1.0,
            "top": self.agent_profile.get("margin_top_inches") if self.agent_profile else 1.0,
            "bottom": self.agent_profile.get("margin_bottom_inches") if self.agent_profile else 1.0,
        }

        if not self.doc.sections:
            return {"success": False, "error": "No sections found in document."}

        actual = {}
        section = self.doc.sections[0]
        actual["left"] = section.left_margin.inches
        actual["right"] = section.right_margin.inches
        actual["top"] = section.top_margin.inches
        actual["bottom"] = section.bottom_margin.inches

        violations = []
        for margin_type, required in requirements.items():
            if abs(actual[margin_type] - required) > 0.05:  # 0.05 inch tolerance
                violations.append({
                    "type": "margin",
                    "margin": margin_type,
                    "required": required,
                    "actual": round(actual[margin_type], 2),
                    "severity": "major",
                })

        return {
            "success": len(violations) == 0,
            "violations": violations,
            "actual": actual,
        }

    def check_font_properties(self) -> Dict[str, Any]:
        """Check font family, size, and style against profile preferences"""
        if not self.doc:
            return {"success": True, "error": "No docx file provided for font checking."}
        
        # Use profile settings if available, otherwise use defaults
        requirements = {
            "font_family": self.agent_profile.get("font_family") if self.agent_profile else "Times New Roman",
            "font_size": self.agent_profile.get("font_size") if self.agent_profile else 12,
            "font_style": self.agent_profile.get("font_style") if self.agent_profile else "normal",
        }

        violations = []

        for para_idx, para in enumerate(self.doc.paragraphs):
            for run in para.runs:
                # Check font family
                if run.font.name and run.font.name != requirements["font_family"]:
                    violations.append({
                        "type": "font_family",
                        "paragraph_index": para_idx,
                        "required": requirements["font_family"],
                        "actual": run.font.name,
                        "severity": "major",
                    })
                    break  # Report once per paragraph

                # Check font size
                if run.font.size:
                    actual_size = run.font.size.pt
                    if actual_size != requirements["font_size"]:
                        violations.append({
                            "type": "font_size",
                            "paragraph_index": para_idx,
                            "required": requirements["font_size"],
                            "actual": actual_size,
                            "severity": "major",
                        })
                        break
                
                # Check font style (bold/italic)
                expected_bold = requirements["font_style"] in ["bold", "bold-italic"]
                expected_italic = requirements["font_style"] in ["italic", "bold-italic"]
                
                if run.font.bold != expected_bold or run.font.italic != expected_italic:
                    violations.append({
                        "type": "font_style",
                        "paragraph_index": para_idx,
                        "required": requirements["font_style"],
                        "actual": f"bold={run.font.bold}, italic={run.font.italic}",
                        "severity": "major",
                    })
                    break

        return {
            "success": len(violations) == 0,
            "violations": violations,
        }

    def check_paragraph_formatting(self) -> Dict[str, Any]:
        """Check line spacing, paragraph spacing, and indentation against profile preferences"""
        if not self.doc:
            return {"success": True, "error": "No docx file provided for paragraph checking."}
        
        # Use profile settings if available
        requirements = {
            "line_spacing": self.agent_profile.get("line_spacing") if self.agent_profile else 2.0,
            "spacing_before": self.agent_profile.get("paragraph_spacing_before") if self.agent_profile else 0,
            "spacing_after": self.agent_profile.get("paragraph_spacing_after") if self.agent_profile else 0,
            "first_line_indent": self.agent_profile.get("first_line_indent") if self.agent_profile else 0.5,
        }

        violations = []
        
        for para_idx, para in enumerate(self.doc.paragraphs):
            # Skip empty paragraphs
            if not para.text.strip():
                continue
            
            # Check line spacing
            if para.paragraph_format.line_spacing:
                actual_spacing = para.paragraph_format.line_spacing
                if abs(actual_spacing - requirements["line_spacing"]) > 0.1:
                    violations.append({
                        "type": "line_spacing",
                        "paragraph_index": para_idx,
                        "required": requirements["line_spacing"],
                        "actual": round(actual_spacing, 2),
                        "severity": "minor",
                    })
            
            # Check first line indent (convert from Pt to inches)
            if para.paragraph_format.first_line_indent:
                actual_indent_inches = para.paragraph_format.first_line_indent.inches
                if abs(actual_indent_inches - requirements["first_line_indent"]) > 0.05:
                    violations.append({
                        "type": "first_line_indent",
                        "paragraph_index": para_idx,
                        "required": requirements["first_line_indent"],
                        "actual": round(actual_indent_inches, 2),
                        "severity": "minor",
                    })
            
        return {
            "success": len(violations) == 0,
            "violations": violations,
        }

    def check_image_properties(self) -> Dict[str, Any]:
        """Check image format and dimensions against profile preferences"""
        if not self.doc:
            return {"success": True, "error": "No docx file provided for image checking."}
        
        # Use profile settings if available
        max_width = self.agent_profile.get("image_max_width_inches") if self.agent_profile else 6.0
        
        violations = []
        image_count = 0
        
        # Iterate through all relationships to find images
        for rel in self.doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
        
        # Check inline shapes (images embedded in runs)
        for para_idx, para in enumerate(self.doc.paragraphs):
            for run in para.runs:
                # Check if run contains an inline shape (image)
                if hasattr(run, '_element') and hasattr(run._element, 'drawing_lst'):
                    for drawing in run._element.drawing_lst:
                        # Check image width if available
                        if hasattr(drawing, 'inline') and drawing.inline:
                            width_emu = drawing.inline.extent.cx  # EMUs (English Metric Units)
                            width_inches = width_emu / 914400  # Convert EMUs to inches
                            
                            if width_inches > max_width:
                                violations.append({
                                    "type": "image_width",
                                    "paragraph_index": para_idx,
                                    "required_max": max_width,
                                    "actual": round(width_inches, 2),
                                    "severity": "minor",
                                })
        
        return {
            "success": len(violations) == 0,
            "violations": violations,
            "images_found": image_count,
        }


    async def evaluate_ai_content(self, document_content: str) -> Dict[str, Any]:
        """
        Evaluates the provided document content using AI.
        Returns structured feedback with severity levels.
        """
        messages = [
            {"role": "system", "content": self.system_prompt},
            {"role": "user", "content": f"Please evaluate the following thesis extract:\n\n{document_content}"}
        ]
        
        response = await self.chat_completion(messages=messages, model=self.agent_model)
        
        # Try to parse as JSON, fall back to structured text response
        try:
            import json
            # Extract JSON from response if it contains code blocks
            if "```json" in response:
                json_str = response.split("```json")[1].split("```")[0].strip()
            elif "```" in response:
                json_str = response.split("```")[1].split("```")[0].strip()
            else:
                json_str = response
            
            return json.loads(json_str)
        except:
            # Return structured response if JSON parsing fails
            return {
                "issues": [],
                "critique": response,
                "parsed_as_text": True
            }

    async def run_analysis(self, document_content: str = "") -> Dict[str, Any]:
        """
        Run full technical analysis combining formatting checks and AI critique.
        Returns structured results suitable for surgical injection.
        """
        # 1. Structural Checks (Pure Python) - Respect enable flags
        all_violations = []
        checks_performed = {}
        
        # Check margins if enabled
        if self.agent_profile.get("enable_margin_check", True):
            margin_check = self.check_margins()
            all_violations.extend(margin_check.get("violations", []))
            checks_performed["margins"] = margin_check.get("success", True)
        else:
            checks_performed["margins"] = None  # Skipped
            
        # Check fonts if enabled
        if self.agent_profile.get("enable_font_check", True):
            font_check = self.check_font_properties()
            all_violations.extend(font_check.get("violations", []))
            checks_performed["fonts"] = font_check.get("success", True)
        else:
            checks_performed["fonts"] = None  # Skipped
            
        # Check paragraphs if enabled
        if self.agent_profile.get("enable_paragraph_check", True):
            paragraph_check = self.check_paragraph_formatting()
            all_violations.extend(paragraph_check.get("violations", []))
            checks_performed["paragraphs"] = paragraph_check.get("success", True)
        else:
            checks_performed["paragraphs"] = None  # Skipped
            
        # Check images if enabled
        if self.agent_profile.get("enable_image_check", True):
            image_check = self.check_image_properties()
            all_violations.extend(image_check.get("violations", []))
            checks_performed["images"] = image_check.get("success", True)
            checks_performed["images_found"] = image_check.get("images_found", 0)
        else:
            checks_performed["images"] = None  # Skipped
            checks_performed["images_found"] = 0

        major_errors = len([v for v in all_violations if v.get("severity") == "major"])
        minor_errors = len([v for v in all_violations if v.get("severity") == "minor"])
        
        # 2. AI Content Review
        ai_feedback = {"issues": [], "critique": "No document content provided for AI evaluation."}
        if document_content.strip():
            ai_feedback = await self.evaluate_ai_content(document_content)

        # Count AI issues by severity (handle both "major"/"minor" and "High"/"Medium"/"Low")
        ai_issues = ai_feedback.get("issues", [])
        ai_major = len([
            i for i in ai_issues 
            if str(i.get("severity", "")).lower() in ["major", "high", "critical"]
        ])
        ai_minor = len([
            i for i in ai_issues 
            if str(i.get("severity", "")).lower() in ["minor", "medium", "low"]
        ])

        # Combine all issues for injection
        all_issues = all_violations + ai_issues
        
        # Calculate overall formatting success (only for enabled checks)
        enabled_checks = [v for v in checks_performed.values() if v is not None]
        formatting_success = all(enabled_checks) if enabled_checks else True
        
        return {
            "agent": "technical_reader",
            "formatting": {
                "success": formatting_success,
                "major_errors": major_errors,
                "minor_errors": minor_errors,
                "violations": all_violations,
                "summary": f"Found {major_errors} major and {minor_errors} minor formatting issues",
                "checks_performed": checks_performed
            },
            "ai_critique": ai_feedback,
            "all_issues": all_issues,
            "major_issues": major_errors + ai_major,
            "minor_issues": minor_errors + ai_minor,
            "total_issues": len(all_issues),
        }