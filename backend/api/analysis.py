"""
Flask Routes for Analysis Execution
Processes documents synchronously for Vercel serverless deployment
Saves original and processed documents to database with surgical injection
"""

from flask import Blueprint, request, jsonify, send_file
from typing import Dict, Any, Optional
from agents.technical_reader import TechnicalReaderAgent
from agents.grammar_critic import GrammarCriticAgent
from agents.statistics import StatisticsAgent
from agents.subject_matter import SubjectMatterExpertAgent
from agents.chairman import ChairmanAgent
from document.surgical_injector import SurgicalInjector
from document.parser import ChapterExtractor
from supabase import create_client
from core.config import settings
from docx import Document
from io import BytesIO
import asyncio
import uuid
import os
import tempfile
from datetime import datetime, timedelta

analysis_bp = Blueprint("analysis", __name__, url_prefix="/api/v1/analysis")

def get_supabase():
    """Get Supabase client connected to the live project"""
    return create_client(settings.SUPABASE_URL, settings.SUPABASE_SERVICE_KEY)


def get_user_id_from_token(token: str) -> str:
    """Extract user ID from JWT token via Supabase"""
    try:
        supabase = get_supabase()
        user = supabase.auth.get_user(token)
        return user.user.id if user and user.user else None
    except Exception as e:
        print(f"Token validation error: {e}")
        return None


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file for LLM analysis."""
    try:
        doc = Document(BytesIO(file_bytes))
        text_content = []
        
        # Extract from paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                text_content.append(p.text)
        
        # Extract from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        
        extracted_text = "\n".join(text_content)
        print(f"[EXTRACT_TEXT] Extracted {len(extracted_text)} chars from document ({len(text_content)} text segments)")
        return extracted_text
    except Exception as e:
        print(f"[EXTRACT_TEXT] Failed to extract DOCX text: {e}")
        import traceback
        traceback.print_exc()
        return ""


def build_formatting_issues(docx_path: str, agent_profile: Optional[Dict[str, Any]], paragraph_texts: Dict[int, str]) -> List[Dict[str, Any]]:
    """Run deterministic formatting checks and return issue objects with locations."""
    profile = agent_profile or {}
    formatting_agent = TechnicalReaderAgent(docx_path=docx_path, agent_profile=profile)
    violations: List[Dict[str, Any]] = []

    if profile.get("enable_margin_check", True):
        margin_check = formatting_agent.check_margins()
        violations.extend(margin_check.get("violations", []))

    if profile.get("enable_font_check", True):
        font_check = formatting_agent.check_font_properties()
        violations.extend(font_check.get("violations", []))

    if profile.get("enable_paragraph_check", True):
        paragraph_check = formatting_agent.check_paragraph_formatting()
        violations.extend(paragraph_check.get("violations", []))

    if profile.get("enable_image_check", True):
        image_check = formatting_agent.check_image_properties()
        violations.extend(image_check.get("violations", []))

    issues: List[Dict[str, Any]] = []
    for violation in violations:
        paragraph_index = violation.get("paragraph_index")
        location_text = ""
        if paragraph_index is not None:
            location_text = paragraph_texts.get(paragraph_index, "")

        if not location_text:
            location_text = f"Formatting issue: {violation.get('type', 'formatting')}"

        issue_text = f"{violation.get('type', 'formatting')} mismatch"
        required = violation.get("required") or violation.get("required_max")
        actual = violation.get("actual")
        if required is not None and actual is not None:
            issue_text = f"{violation.get('type', 'formatting')} mismatch (expected {required}, found {actual})"

        suggestion = "Update formatting to match the profile requirements."
        if required is not None:
            suggestion = f"Set to {required} as required by the profile."

        issues.append({
            "type": violation.get("type", "formatting"),
            "location": {"text": location_text},
            "severity": violation.get("severity", "major"),
            "issue": issue_text,
            "suggestion": suggestion,
            "source": "formatting_check",
        })

    return issues


def run_agent_analysis(agent_role: str, document_text: str, parser: ChapterExtractor, agent_profile: Optional[Dict[str, Any]], max_chunks: int = 50) -> Dict[str, Any]:
    """Run agent analysis with chunking for large documents."""
    total_chars = len(document_text)
    if total_chars > 15000:
        chunks = parser.chunk_for_llm(max_tokens=8000)
        max_chunks_to_analyze = min(len(chunks), max_chunks)
        all_chunk_issues: List[Dict[str, Any]] = []

        for chunk_idx, chunk in enumerate(chunks[:max_chunks_to_analyze]):
            print(f"[ANALYSIS] Processing {agent_role} chunk {chunk_idx + 1}/{max_chunks_to_analyze}")
            agent = get_agent_by_role(agent_role, agent_profile)
            chunk_result = asyncio.run(agent.run_analysis(chunk['text']))

            if isinstance(chunk_result, dict) and 'issues' in chunk_result:
                all_chunk_issues.extend(chunk_result['issues'])

        return {
            "issues": all_chunk_issues,
            "chunked_analysis": True,
            "chunks_processed": max_chunks_to_analyze,
            "total_chunks": len(chunks),
        }

    agent = get_agent_by_role(agent_role, agent_profile)
    result = asyncio.run(agent.run_analysis(document_text))
    return result if isinstance(result, dict) else {"error": str(result)}


def get_agent_by_role(agent_role: str, agent_profile: Optional[Dict[str, Any]] = None):
    """Factory function to get the appropriate agent based on role"""
    agent_map = {
        "tech": TechnicalReaderAgent,
        "grammar": GrammarCriticAgent,
        "stats": StatisticsAgent,
        "subject": SubjectMatterExpertAgent,
        "chairman": ChairmanAgent,
    }
    
    agent_class = agent_map.get(agent_role)
    if not agent_class:
        raise ValueError(f"Unknown agent role: {agent_role}")
    
    return agent_class(agent_profile=agent_profile)


@analysis_bp.route("/start", methods=["POST"])
def start_analysis():
    """
    Start analysis for a given file and agent role.
    1. Downloads original document
    2. Runs selected agent analysis
    3. Applies surgical injection with highlights
    4. Saves both original and processed documents
    5. Returns task info with processed file path
    """
    # Get user from JWT token
    auth_header = request.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        return {"error": "Unauthorized"}, 401

    token = auth_header.replace("Bearer ", "")
    user_id = get_user_id_from_token(token)

    if not user_id:
        return {"error": "Invalid token"}, 401

    try:
        data = request.get_json()
        file_id = data.get("file_id")
        agent_role = data.get("agent_role")
        profile_id = data.get("profile_id")  # Optional: use specific profile
        
        if not file_id or not agent_role:
            return {"error": "Missing file_id or agent_role"}, 400

        # Validate agent role
        valid_roles = ["tech", "grammar", "stats", "subject", "chairman"]
        if agent_role not in valid_roles:
            return {"error": f"Invalid agent_role. Must be one of: {', '.join(valid_roles)}"}, 400

        supabase = get_supabase()
        
        # Retrieve document from Supabase
        doc_record = (
            supabase.table("documents")
            .select("*")
            .eq("id", file_id)
            .eq("owner_id", user_id)
            .single()
            .execute()
        )
        
        if not doc_record.data:
            return {"error": "Document not found"}, 404

        # Download file from Supabase Storage
        file_path = doc_record.data["file_path"]
        original_file_bytes = supabase.storage.from_(settings.STORAGE_BUCKET_NAME).download(file_path)

        # Generate task ID for tracking
        task_id = str(uuid.uuid4())

        # Get agent profile (prioritize profile_id if provided, then use active profile, then use defaults)
        agent_profile_record = None
        try:
            if profile_id:
                # Use specific profile if provided (user-owned or system default)
                profile_result = (
                    supabase.table("agent_profiles")
                    .select("*")
                    .eq("id", profile_id)
                    .eq("agent_role", agent_role)
                    .or_(f"owner_id.eq.{user_id},owner_id.is.null")
                    .single()
                    .execute()
                )
                if profile_result.data:
                    agent_profile_record = profile_result.data
            else:
                # Use active profile for this role
                profiles = (
                    supabase.table("agent_profiles")
                    .select("*")
                    .eq("owner_id", user_id)
                    .eq("agent_role", agent_role)
                    .eq("is_active", True)
                    .execute()
                )
                if profiles.data and len(profiles.data) > 0:
                    agent_profile_record = profiles.data[0]
        except Exception as e:
            print(f"Warning: Could not fetch agent profile: {e}")

        # Run analysis
        analysis_result: Dict[str, Any] = {}
        processed_file_bytes = None
        error_message = None

        try:
            # Save original file to temp location first for parsing
            temp_original = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            temp_original.write(original_file_bytes)
            temp_original.close()
            
            # Extract document structure for frontend display
            parser = ChapterExtractor(temp_original.name)
            document_structure = parser.extract_full_structure(max_paragraphs=10000)
            
            # Validate document size (200 pages max ≈ 50,000 words)
            estimated_pages = document_structure.get('estimated_pages', 0)
            if estimated_pages > 200:
                error_message = f"Document too large: {estimated_pages} estimated pages (max: 200 pages)"
                raise ValueError(error_message)
            
            print(f"[ANALYSIS] Extracted document structure: {document_structure['extracted_paragraphs']} paragraphs, ~{estimated_pages} pages")
            
            # Extract text from document for AI analysis
            document_text = extract_text_from_docx(original_file_bytes)
            total_chars = len(document_text)
            print(f"[ANALYSIS] Extracted {total_chars} chars from document")
            
            # Run main agent analysis (with chunking for large documents)
            print(f"[ANALYSIS] Running {agent_role} agent analysis...")
            analysis_result = run_agent_analysis(
                agent_role,
                document_text,
                parser,
                agent_profile_record,
                max_chunks=50,
            )

            # Always run grammar/spelling checks (even for non-grammar roles)
            if agent_role != "grammar":
                print("[ANALYSIS] Running mandatory grammar/spelling checks...")
                grammar_result = run_agent_analysis(
                    "grammar",
                    document_text,
                    parser,
                    agent_profile_record,
                    max_chunks=50,
                )
                grammar_issues = grammar_result.get("issues", []) if isinstance(grammar_result, dict) else []
                if grammar_issues:
                    existing_issues = analysis_result.get("issues", [])
                    if not isinstance(existing_issues, list):
                        existing_issues = []
                    analysis_result["issues"] = existing_issues + grammar_issues
                    analysis_result["total_issues"] = len(analysis_result["issues"])
                    analysis_result["major_errors"] = len([
                        i for i in analysis_result["issues"]
                        if str(i.get("severity", "")).lower() in ["major", "high", "critical"]
                    ])
                    analysis_result["minor_errors"] = len([
                        i for i in analysis_result["issues"]
                        if str(i.get("severity", "")).lower() in ["minor", "medium", "low"]
                    ])
                analysis_result["grammar_check"] = grammar_result
            
            print(f"[ANALYSIS] Agent returned result with keys: {list(analysis_result.keys())}")
            print(f"[ANALYSIS] Issues count: {len(analysis_result.get('issues', []))}")
            
            # Add document structure to analysis results
            analysis_result["document_structure"] = document_structure

            # Run deterministic formatting checks for ALL roles
            paragraph_texts = {
                p.get("index"): p.get("text", "")
                for p in document_structure.get("paragraphs", [])
                if isinstance(p, dict)
            }
            formatting_issues = build_formatting_issues(
                temp_original.name,
                agent_profile_record,
                paragraph_texts,
            )
            if formatting_issues:
                existing_issues = analysis_result.get("issues")
                if not isinstance(existing_issues, list) or not existing_issues:
                    existing_issues = analysis_result.get("all_issues", [])
                if not isinstance(existing_issues, list):
                    existing_issues = []
                analysis_result["issues"] = existing_issues + formatting_issues
                analysis_result["formatting_issues"] = formatting_issues
                analysis_result["total_issues"] = len(analysis_result["issues"])
                analysis_result["major_errors"] = len([
                    i for i in analysis_result["issues"]
                    if str(i.get("severity", "")).lower() in ["major", "high", "critical"]
                ])
                analysis_result["minor_errors"] = len([
                    i for i in analysis_result["issues"]
                    if str(i.get("severity", "")).lower() in ["minor", "medium", "low"]
                ])
            
            # Normalize agent response to include all_issues with agent info
            if "issues" in analysis_result:
                # Transform issues array and add agent info to each issue
                normalized_issues = []
                for issue in analysis_result.get("issues", []):
                    if isinstance(issue, dict):
                        normalized_issue = issue.copy()
                        normalized_issue["agent"] = agent_role
                        normalized_issues.append(normalized_issue)
                
                analysis_result["all_issues"] = normalized_issues
                print(f"[ANALYSIS] Normalized {len(normalized_issues)} issues to all_issues")
            
            # Also check for alternative field names used by different agents
            if not analysis_result.get("all_issues"):
                for key in ["issues", "findings", "problems", "errors"]:
                    if key in analysis_result and isinstance(analysis_result[key], list):
                        normalized_issues = []
                        for issue in analysis_result[key]:
                            if isinstance(issue, dict):
                                normalized_issue = issue.copy()
                                normalized_issue["agent"] = agent_role
                                normalized_issues.append(normalized_issue)
                        analysis_result["all_issues"] = normalized_issues
                        print(f"[ANALYSIS] Normalized {len(normalized_issues)} from '{key}' to all_issues")
                        break

            # Apply surgical injection to create processed document
            # temp_original already created above for parsing
            try:
                # Create surgical injector
                injector = SurgicalInjector(temp_original.name)
                
                # Apply highlights based on all_issues if available
                all_issues = analysis_result.get("all_issues", [])
                if all_issues:
                    inject_result = injector.inject_issue_highlights(all_issues)
                    analysis_result["injection_summary"] = inject_result
                
                # Save the processed document
                temp_processed = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
                temp_processed.close()  # Close the file so we can write to it
                injector.save_processed_file(temp_processed.name)
                
                # Read the processed file bytes
                with open(temp_processed.name, 'rb') as f:
                    processed_file_bytes = f.read()
                
                # Clean up temp file
                try:
                    os.unlink(temp_processed.name)
                except:
                    pass
                    
            finally:
                # Clean up original temp file
                try:
                    os.unlink(temp_original.name)
                except:
                    pass

        except Exception as agent_error:
            error_message = str(agent_error)
            print(f"Analysis error: {error_message}")
            result = None

        # Extract major/minor errors for fast access
        major_errors = analysis_result.get("major_errors", 0) or analysis_result.get("major_issues", 0)
        minor_errors = analysis_result.get("minor_errors", 0) or analysis_result.get("minor_issues", 0)
        total_issues = analysis_result.get("total_issues", 0)

        # Create analysis task record
        analysis_task = {
            "id": task_id,
            "owner_id": user_id,
            "file_id": file_id,
            "agent_role": agent_role,
            "status": "failed" if error_message else "complete",
            "progress": 100 if not error_message else 0,
            "results": analysis_result,
            "error_message": error_message,
            "created_at": datetime.utcnow().isoformat(),
            "completed_at": datetime.utcnow().isoformat(),
        }

        # Save analysis task to database
        try:
            supabase.table("analysis_tasks").insert(analysis_task).execute()
        except Exception as e:
            print(f"Error saving analysis task: {e}")

        # If processing was successful, save processed document
        if not error_message and processed_file_bytes:
            try:
                # Generate path for processed file
                processed_file_id = str(uuid.uuid4())
                original_name = os.path.splitext(doc_record.data["name"])[0]
                processed_filename = f"{original_name}_analyzed_{agent_role}.docx"
                processed_file_path = f"{user_id}/{processed_file_id}/{processed_filename}"

                # Upload processed file to storage
                supabase.storage.from_(settings.STORAGE_BUCKET_NAME).upload(
                    processed_file_path,
                    processed_file_bytes,
                    {"cacheControl": "3600", "upsert": "false"},
                )

                # Create record for processed document
                processed_doc_record = {
                    "id": processed_file_id,
                    "owner_id": user_id,
                    "name": processed_filename,
                    "file_path": processed_file_path,
                    "status": "analyzed",
                    "major_errors": major_errors,
                    "minor_errors": minor_errors,
                    "created_at": datetime.utcnow().isoformat(),
                    "expired_at": (datetime.utcnow() + timedelta(hours=1)).isoformat(),
                }

                supabase.table("documents").insert(processed_doc_record).execute()

                # Update original document status
                supabase.table("documents").update({
                    "status": "analyzed",
                    "major_errors": major_errors,
                    "minor_errors": minor_errors,
                }).eq("id", file_id).execute()

                # Update analysis task with processed file info
                analysis_task_results = analysis_result.copy()
                analysis_task_results["processed_file_path"] = processed_file_path
                analysis_task_results["processed_file_id"] = processed_file_id
                analysis_task_results["major_errors"] = major_errors
                analysis_task_results["minor_errors"] = minor_errors
                
                supabase.table("analysis_tasks").update({
                    "results": analysis_task_results,
                }).eq("id", task_id).execute()

                return {
                    "task_id": task_id,
                    "file_id": file_id,
                    "processed_file_id": processed_file_id,
                    "agent_role": agent_role,
                    "status": "complete",
                    "major_errors": major_errors,
                    "minor_errors": minor_errors,
                    "total_issues": total_issues,
                    "result": analysis_task_results,
                    "original_file_path": file_path,
                    "processed_file_path": processed_file_path,
                }, 200

            except Exception as e:
                error_msg = f"Storage error: {str(e)}"
                print(error_msg)
                try:
                    supabase.table("analysis_tasks").update({
                        "status": "failed",
                        "error_message": error_msg,
                    }).eq("id", task_id).execute()
                except Exception as update_error:
                    print(f"Error updating analysis task failure: {update_error}")

                return {
                    "task_id": task_id,
                    "file_id": file_id,
                    "agent_role": agent_role,
                    "status": "failed",
                    "error": error_msg,
                }, 500

        # Return error if analysis failed
        if error_message:
            return {
                "task_id": task_id,
                "file_id": file_id,
                "agent_role": agent_role,
                "status": "failed",
                "error": error_message,
            }, 500

        return {
            "task_id": task_id,
            "file_id": file_id,
            "agent_role": agent_role,
            "status": "complete",
            "major_errors": major_errors,
            "minor_errors": minor_errors,
            "total_issues": total_issues,
            "result": analysis_result,
        }, 200

    except Exception as e:
        return {"error": f"Analysis failed: {str(e)}"}, 500


@analysis_bp.route("/status/<task_id>", methods=["GET"])
def get_analysis_status(task_id):
    """
    Get the status of an analysis task.
    """
    auth_header = request.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        return {"error": "Unauthorized"}, 401

    token = auth_header.replace("Bearer ", "")
    user_id = get_user_id_from_token(token)

    if not user_id:
        return {"error": "Invalid token"}, 401

    try:
        supabase = get_supabase()
        task = (
            supabase.table("analysis_tasks")
            .select("*")
            .eq("id", task_id)
            .eq("owner_id", user_id)
            .single()
            .execute()
        )

        if not task.data:
            return {"error": "Task not found"}, 404

        return jsonify(task.data), 200

    except Exception as e:
        return {"error": str(e)}, 500


@analysis_bp.route("/download/<task_id>", methods=["GET"])
def download_processed_file(task_id):
    """
    Download the processed DOCX file after analysis.
    """
    auth_header = request.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        return {"error": "Unauthorized"}, 401

    token = auth_header.replace("Bearer ", "")
    user_id = get_user_id_from_token(token)

    if not user_id:
        return {"error": "Invalid token"}, 401

    try:
        supabase = get_supabase()

        # Get analysis task
        task = (
            supabase.table("analysis_tasks")
            .select("*")
            .eq("id", task_id)
            .eq("owner_id", user_id)
            .single()
            .execute()
        )

        if not task.data:
            return {"error": "Task not found"}, 404

        # Get processed document info from results
        results = task.data.get("results", {})
        if isinstance(results, str):
            import json
            results = json.loads(results)

        processed_file_path = results.get("processed_file_path")
        if not processed_file_path:
            return {"error": "Processed file not found"}, 404

        # Download file from storage
        file_data = supabase.storage.from_(settings.STORAGE_BUCKET_NAME).download(processed_file_path)

        return send_file(
            BytesIO(file_data),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="analyzed_document.docx",
        )

    except Exception as e:
        return {"error": str(e)}, 500
